# nazwa pliku powinna być create_declaration.py

from docx import Document
from CountParkingFee import CountParkingFee

 # czy to dziedziczenie jest tutaj potrzebne, tylko po to, by mieć dostęp do 
 # self.count_days i tych różnych self.fees? bardzo słaby pomysł,
# bo Proces tworzenie deklaracji ma bardzo mało wspólnego z opłatą
# Idealnie, to API wyglądało by tak:
# place = ParkingPlace(home_port, chip_card, parking_place)
# yaht = Yaht(name, registration, withth, length, type)
# owner = Owner(name, ...)
# declaration = Declaration(place = place, yaht = yaht, owner = owner)
# declaration.save('dupa.docx')

class CreateDeclaration(CountParkingFee):
    
    def __init__(self, place: ParkingPlace, yaht: Yaht, owner: Person):
        super().__init__(parking_peroid, yacht_length, yacht_width)
        self.document = Document('deklaracja.docx')
        self.parking_peroid = parking_peroid
        self.days = self.count_days()
        self.parking_place = parking_place
        self.date = date
        # ---- stwórz klase Yacht z tymi parametrami:
        # albo zrób tak samo jak z owner_details
        self.name_yacht = name_yacht
        self.registration_number = registration_number
        self.home_port = home_port
        self.yacht_length = yacht_length
        self.yacht_width = yacht_width
        self.yacht_type = yacht_type
        # ----
        self.owner_details = owner_details
        self.commissioning_body = commissioning_body

        self.parking_fee = round(self.parking_fee(), 2)
        self.quarter_fee = round(self.parking_fee / 4.0, 2)
        self.chip_card = chip_card

    def create_document(self):
        for paragraph in self.document.paragraphs:
            if 'Miejsce postojowe' in paragraph.text:
                paragraph.text = 'Miejsce postojowe {}                                                                  ' \
                                 '                                        ' \
                                 'Data {}'.format(self.parking_place, self.date)

        self.document.add_paragraph('\n')
        self.document.add_paragraph('Nazwa jachtu: ', style='List Number').add_run(self.name_yacht).bold = True
        self.document.add_paragraph('Nr rejestracyjny: ', style='List Number').add_run(
            self.registration_number).bold = True
        self.document.add_paragraph('Port macierzysty: ', style='List Number').add_run(self.home_port).bold = True
        self.document.add_paragraph('Dane jachtu: ', style='List Number')
        p = self.document.add_paragraph('            Długość: ')
        p.add_run(str(self.yacht_length)).bold = True
        p.add_run('             Szerokość: ')
        p.add_run(str(self.yacht_width)).bold = True
        self.document.add_paragraph('Typ jachtu: ', style='List Number').add_run(self.yacht_type).bold = True
        self.document.add_paragraph('Dane właściciela* lub użytkownika jachtu*: ', style='List Number')
        self.document.add_paragraph('            - imię i nazwisko: ').add_run(self.owner_details['name']).bold = True
        self.document.add_paragraph('            - adres : ').add_run(self.owner_details['address']).bold = True
        self.document.add_paragraph('Dane armatora: ', style='List Number')
        self.document.add_paragraph('            - imię i nazwisko: ').add_run(self.owner_details['name']).bold = True
        self.document.add_paragraph('            - adres : ').add_run(self.owner_details['address']).bold = True
        self.document.add_paragraph('Podmiot zlecający, podpisujący umowę na postój jachtu: ', style='List Number')
        self.document.add_paragraph('            - nazwisko imię */ pełna nazwa klubu lub firmy: ').add_run(
            self.owner_details['name']).bold = True
        self.document.add_paragraph('            - adres: ').add_run(self.commissioning_body['address']).bold = True
        self.document.add_paragraph('            - tel: ').add_run(self.commissioning_body['tel']).bold = True
        self.document.add_paragraph('            - E-mial: ').add_run(self.commissioning_body['e-mail']).bold = True
        self.document.add_paragraph('            - NIP klubu/stowarzyszenia: ').add_run(
            self.commissioning_body['nip']).bold = True
        self.document.add_paragraph('Deklarowany okres i czas postoju: ', style='List Number')
        self.document.add_paragraph('            - od dnia: ').add_run(self.parking_peroid['from']).bold = True
        self.document.add_paragraph('            - do dnia: ').add_run(self.parking_peroid['to']).bold = True
        self.document.add_paragraph('')
        self.document.add_paragraph('Opłata za postój wynosi: ', style='List Number').add_run(
            str(self.parking_fee)).bold = True
        self.document.add_paragraph(
            'Czynsz płatny, zgodnie z wystawioną fakturą z góry, jednorazowo lub w czterech poniższych '
            'ratach za każdy kwartał do:')
        self.document.add_paragraph('      I. 15.05.2020, w kwocie: ').add_run(str(self.quarter_fee)).bold = True
        self.document.add_paragraph('     II. 15.08.2020, w kwocie: ').add_run(str(self.quarter_fee)).bold = True
        self.document.add_paragraph('    III. 15.11.2020, w kwocie: ').add_run(str(self.quarter_fee)).bold = True
        self.document.add_paragraph('    IV. 15.21.2020, w kwocie: ').add_run(str(self.quarter_fee)).bold = True
        self.document.add_paragraph('na rachunek Wynajmującego o nr  88 1030 1117 0000 0000 8899 5007.')
        self.document.add_paragraph('Adres do korespondencji: ', style='List Number').add_run(
            self.owner_details['address']).bold = True
        self.document.add_paragraph('Karta chipowa: ', style='List Number').add_run(self.chip_card).bold = True
        self.document.add_paragraph('Miejsce postoju nr: ', style='List Number').add_run(self.parking_place).bold = True
        self.document.add_paragraph(
            '\nUprzejmie informujemy, że w trakcie postoju jachtu na przystani NCŻ AWFiS może zaistnieć '
            'konieczność zmiany miejsca lokalizacji postoju jachtu na inne zgodnie z ze wskazaniem '
            'upoważnionego pracownika przystani.\n\n'
            'Oświadczam, iż zapoznałem się z Regulaminem przystani Narodowego Centrum Żeglarstwa AWFiS '
            'Gdańsku, w tym zawarcia umowy na postój. W pełni go akceptuję co poświadczam własnoręcznym '
            'podpisem pod Deklaracją postoju. Oświadczam, iż w przypadku okresu postoju krótszego niż '
            'zadeklarowany (wynikający m.in. ze sprzedaży jachtu) mam świadomość, że stawka będzie '
            'rekalkulowana do krótszego okresu postoju, zgodnie z obowiązującym cennikiem.\n\n').paragraph_format.alignment = 3
        self.document.add_paragraph(".........................................................				 .........."
                                    "...............................................").add_run(
            '   podpis pracownika NCŻ                                                                                 '
            '             podpis (imię, nazwisko)').italic = True

        self.document.add_page_break()

        self.document.save('Deklaracja {}.docx'.format(self.name_yacht))
